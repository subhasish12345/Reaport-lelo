"""
report_generator.py
-------------------
Core docx generation logic — imported by both the CLI script and the web app.
Call generate_report_bytes(content: str) -> bytes to get the .docx as bytes.
"""
import io
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _create_element(name):
    return OxmlElement(name)

def _create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def _add_page_number(run):
    fldChar1 = _create_element('w:fldChar')
    _create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = _create_element('w:instrText')
    _create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    fldChar2 = _create_element('w:fldChar')
    _create_attribute(fldChar2, 'w:fldCharType', 'separate')
    fldChar3 = _create_element('w:fldChar')
    _create_attribute(fldChar3, 'w:fldCharType', 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def _add_pageref(paragraph, bookmark_name):
    """Insert a Word field (PAGEREF) into a paragraph."""
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(ns.qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(ns.qn('xml:space'), 'preserve')
    instrText.text = f' PAGEREF {bookmark_name} \\h '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(ns.qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(ns.qn('w:fldCharType'), 'end')
    r = run._r
    r.append(fldChar)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def _add_bookmark(paragraph, bookmark_name, bookmark_id):
    """Wrap paragraph content in a bookmark."""
    p = paragraph._p
    
    # 1. Start
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(ns.qn('w:id'), str(bookmark_id))
    bm_start.set(ns.qn('w:name'), bookmark_name)
    
    # 2. End
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(ns.qn('w:id'), str(bookmark_id))
    
    # Insert start at the beginning of the paragraph
    p.insert(0, bm_start)
    # Insert end at the end of the paragraph
    p.append(bm_end)

def _set_cell_border(cell, **kwargs):
    """
    Apply borders to a cell.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Check if tcBorders already exists
    tcBorders = tcPr.find(ns.qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for side, params in kwargs.items():
        tag = 'w:' + side
        border = OxmlElement(tag)
        for key, value in params.items():
            border.set(ns.qn('w:' + key), str(value))
        tcBorders.append(border)

def _style_grid_table(table):
    """Apply borders to all cells in the table."""
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell, 
                top={"sz": 4, "val": "single", "color": "000000"},
                bottom={"sz": 4, "val": "single", "color": "000000"},
                left={"sz": 4, "val": "single", "color": "000000"},
                right={"sz": 4, "val": "single", "color": "000000"}
            )


# ---------------------------------------------------------------------------
# Main generator
# ---------------------------------------------------------------------------

def generate_report_bytes(content: str) -> bytes:
    """
    Parse `content` (formatted text) and return the generated .docx as bytes.
    """
    doc = Document()

    # -----------------------------------------------------------------------
    # PASS 1: Discovery (Collecting TOC and LOF items)
    # -----------------------------------------------------------------------
    toc_items = []
    lof_items = []
    
    bm_counter = 0
    lines = content.split('\n')
    
    for raw_line in lines:
        line = raw_line.strip()
        if not line: continue
        
        if line.startswith('CHAPTER:'):
            text = line.replace('CHAPTER:', '').strip()
            # Special case for Chapter numbering in text
            if text.lower().startswith('chapter '):
                parts = text.split(' ', 2)
                if len(parts) == 3 and parts[1].isdigit():
                    text = f"Chapter {parts[1]}: {parts[2]}"
            
            bm_counter += 1
            toc_items.append({'text': text, 'bookmark': f'bm_h{bm_counter}', 'level': 1})
            
        elif line.startswith('HEADING:'):
            bm_counter += 1
            toc_items.append({'text': line.replace('HEADING:', '').strip(), 'bookmark': f'bm_h{bm_counter}', 'level': 2})
            
        elif line.startswith('SUBHEADING:'):
            bm_counter += 1
            toc_items.append({'text': line.replace('SUBHEADING:', '').strip(), 'bookmark': f'bm_h{bm_counter}', 'level': 3})
            
        elif line.startswith('FIGURE:'):
            bm_counter += 1
            lof_items.append({'text': line.replace('FIGURE:', '').strip(), 'bookmark': f'bm_f{bm_counter}'})


    # -----------------------------------------------------------------------
    # Document Setup
    # -----------------------------------------------------------------------
    # Margins — 1 inch all sides
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # --- Setup Front Matter Numbering (I, II, III) ---
    first_section = doc.sections[0]
    first_section.footer_distance = Inches(0.5)
    footer = first_section.footer
    if not footer.paragraphs: footer.add_paragraph()
    fp = footer.paragraphs[0]
    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fr = fp.add_run()
    fr.font.name = 'Times New Roman'
    fr.font.size = Pt(12)
    _add_page_number(fr)

    # Set numbering format to upperRoman for first section and START AT 4 (IV)
    # (Since Title, Certificate, Ack, Abstract = pages 1-4)
    sectPr = first_section._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(ns.qn('w:fmt'), 'upperRoman')
    pgNumType.set(ns.qn('w:start'), '4')
    sectPr.append(pgNumType)

    styles = doc.styles

    # Normal style — 12pt, double spacing, justified
    ns_style = styles['Normal']
    ns_style.font.name = 'Times New Roman'
    ns_style.font.size = Pt(12)
    ns_style.paragraph_format.line_spacing      = 2.0
    ns_style.paragraph_format.line_spacing_rule = 1
    ns_style.paragraph_format.alignment         = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Heading 1 — Chapter names: 18pt Bold
    h1 = styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment   = WD_PARAGRAPH_ALIGNMENT.LEFT
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after  = Pt(12)

    # Heading 2 — Subheadings: 16pt Bold
    h2 = styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after  = Pt(12)

    # Heading 3 — Sub-subheadings: 14pt Bold
    try:
        h3 = styles['Heading 3']
    except KeyError:
        from docx.enum.style import WD_STYLE_TYPE
        h3 = styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after  = Pt(10)

    # --- Ensure 'Caption' style exists — 10pt italic centered ---
    try:
        cap_style = styles['Caption']
        cap_style.font.name = 'Times New Roman'
        cap_style.font.size = Pt(10)
        cap_style.font.italic = True
        cap_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cap_style.paragraph_format.space_after = Pt(12)
    except KeyError:
        from docx.enum.style import WD_STYLE_TYPE
        cap_style = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
        cap_style.base_style = styles['Normal']
        cap_style.font.italic = True
        cap_style.font.size = Pt(10)
        cap_style.font.name = 'Times New Roman'
        cap_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cap_style.paragraph_format.space_after = Pt(12)

    # -----------------------------------------------------------------------
    # Front matter placeholder
    # -----------------------------------------------------------------------
    doc.add_heading("PROJECT REPORT", level=1)
    doc.add_paragraph("Manual Pages: Title Page, Certificate, Acknowledgement, Abstract")
    doc.add_page_break()

    # -----------------------------------------------------------------------
    # GRID TABLE TOC
    # -----------------------------------------------------------------------
    doc.add_paragraph("Table of Contents", style='Heading 1')
    
    toc_table = doc.add_table(rows=1, cols=2)
    toc_table.width = Inches(6.5) # Full page width
    
    # TOC Headers
    hdr_cells = toc_table.rows[0].cells
    hdr_cells[0].text = "Content"
    hdr_cells[1].text = "Page"
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
    for item in toc_items:
        row = toc_table.add_row().cells
        p0 = row[0].paragraphs[0]
        p0.text = item['text']
        
        # Indentation for subheadings
        if item['level'] == 2: p0.paragraph_format.left_indent = Pt(12)
        if item['level'] == 3: p0.paragraph_format.left_indent = Pt(24)
        if item['level'] == 1: p0.runs[0].bold = True
            
        p1 = row[1].paragraphs[0]
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p1.text = "..." # Placeholder until update
        _add_pageref(p1, item['bookmark'])
    
    _style_grid_table(toc_table)
    doc.add_paragraph("[Select table and press F9 to update page numbers]").italic = True
    doc.add_page_break()

    # -----------------------------------------------------------------------
    # GRID TABLE LOF
    # -----------------------------------------------------------------------
    if lof_items:
        doc.add_paragraph("List of Figures", style='Heading 1')
        lof_table = doc.add_table(rows=1, cols=2)
        
        hdr = lof_table.rows[0].cells
        hdr[0].text = "Figure Description"
        hdr[1].text = "Page"
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
        for item in lof_items:
            row = lof_table.add_row().cells
            row[0].text = item['text']
            p1 = row[1].paragraphs[0]
            p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p1.text = "..." # Placeholder until update
            _add_pageref(p1, item['bookmark'])
            
        _style_grid_table(lof_table)
        doc.add_page_break()

    # -----------------------------------------------------------------------
    # MAIN CONTENT (PASS 2 - Generation)
    # -----------------------------------------------------------------------
    toc_idx = 0
    lof_idx = 0
    bookmark_id = 0
    
    for line in content.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('===') or line.startswith('NOTE') or line.startswith('FakeShield'): continue

        if line.startswith('CHAPTER:'):
            item = toc_items[toc_idx]
            toc_idx += 1
            
            if item['text'].lower().startswith('chapter 1:'):
                new_section = doc.add_section(2)
                new_section.footer.is_linked_to_previous = False
                fp = new_section.footer.paragraphs[0]
                fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                fr = fp.add_run()
                _add_page_number(fr)
                
                # Restart numbering at 1 (decimal) and ensure centered
                for p in new_section.footer.paragraphs:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                sectPr = new_section._sectPr
                pgNumType = OxmlElement('w:pgNumType')
                pgNumType.set(ns.qn('w:start'), "1")
                pgNumType.set(ns.qn('w:fmt'), 'decimal')
                sectPr.append(pgNumType)
            else:
                doc.add_page_break()

            h = doc.add_heading(item['text'], level=1)
            bookmark_id += 1
            _add_bookmark(h, item['bookmark'], bookmark_id)

        elif line.startswith('HEADING:'):
            item = toc_items[toc_idx]
            toc_idx += 1
            h = doc.add_heading(item['text'], level=2)
            bookmark_id += 1
            _add_bookmark(h, item['bookmark'], bookmark_id)

        elif line.startswith('SUBHEADING:'):
            item = toc_items[toc_idx]
            toc_idx += 1
            h = doc.add_heading(item['text'], level=3)
            bookmark_id += 1
            _add_bookmark(h, item['bookmark'], bookmark_id)

        elif line.startswith('PARA:'):
            doc.add_paragraph(line.replace('PARA:', '').strip(), style='Normal')

        elif line.startswith('FIGURE:'):
            item = lof_items[lof_idx]
            lof_idx += 1
            
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Inches(3.5)
            spacer.paragraph_format.keep_with_next = True

            cap = doc.add_paragraph(item['text'], style='Caption')
            # Force caption font/size directly in case style is overridden
            for run in cap.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.italic = True
            cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            bookmark_id += 1
            _add_bookmark(cap, item['bookmark'], bookmark_id)

        else:
            doc.add_paragraph(line)

    # Save
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
